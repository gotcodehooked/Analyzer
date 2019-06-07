import re

from module.DataMiner import DataMiner
from docx import *



class DocxDataMiner(DataMiner):

    def __init__(self, path):
        super().__init__(path)

    def parseData(self):
        pharList = []
        document = Document(self.getpath)
        head1s = []

        for paragraph in document:
            print(paragraph.text)

        # print(head1s)

        # for i in range(len(document.paragraphs)):
        #     if str(document.paragraphs[i].text).startswith("Ключевые "):
        #         pharList.append(document.paragraphs[i].text)
        # for i in range(len(document.paragraphs)):
        #     text = document.paragraphs[i].element
        #     print(text)
        #
        #     yield text

        # return pharList



    def analyzeData(self):
        pharList = self.parseData()
        # print(type(pharList))
        print(pharList)
        # i = 0
        # for i in pharList:
        #     print(i, "\n")
        #     print(i)
        #     i+=1
        # phar = self.parseData()
        # pharList = [pharList for pharList in next(phar)]
        # print(len(pharList))
        # for i in range(len(pharList)):
        #     print(pharList[i])




















































    def sendReport(self):
        pass
